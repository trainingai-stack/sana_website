"""
Markdown to Word Document Converter
将Markdown文件转换为美观的Word文档
"""

import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)


def create_styled_document():
    """创建带样式的文档"""
    doc = Document()
    
    # 设置页面边距
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    # 设置标题样式
    styles = doc.styles
    
    # 标题1样式
    h1_style = styles['Heading 1']
    h1_style.font.name = 'Microsoft YaHei'
    h1_style.font.size = Pt(18)
    h1_style.font.bold = True
    h1_style.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
    h1_style.paragraph_format.space_before = Pt(24)
    h1_style.paragraph_format.space_after = Pt(12)
    
    # 标题2样式
    h2_style = styles['Heading 2']
    h2_style.font.name = 'Microsoft YaHei'
    h2_style.font.size = Pt(16)
    h2_style.font.bold = True
    h2_style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    h2_style.paragraph_format.space_before = Pt(18)
    h2_style.paragraph_format.space_after = Pt(8)
    
    # 标题3样式
    h3_style = styles['Heading 3']
    h3_style.font.name = 'Microsoft YaHei'
    h3_style.font.size = Pt(14)
    h3_style.font.bold = True
    h3_style.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    h3_style.paragraph_format.space_before = Pt(12)
    h3_style.paragraph_format.space_after = Pt(6)
    
    # 正文样式
    normal_style = styles['Normal']
    normal_style.font.name = 'Microsoft YaHei'
    normal_style.font.size = Pt(11)
    normal_style.paragraph_format.line_spacing = 1.5
    normal_style.paragraph_format.space_after = Pt(6)
    
    return doc


def parse_table(lines, start_idx):
    """解析Markdown表格"""
    rows = []
    idx = start_idx
    
    while idx < len(lines) and '|' in lines[idx]:
        row = lines[idx].strip()
        if row.startswith('|'):
            row = row[1:]
        if row.endswith('|'):
            row = row[:-1]
        cells = [cell.strip() for cell in row.split('|')]
        rows.append(cells)
        idx += 1
    
    return rows, idx


def add_table(doc, rows):
    """添加带样式的表格"""
    if not rows:
        return
    
    # 跳过分隔行（包含---的行）
    data_rows = [r for r in rows if not all(c.replace('-', '').replace(':', '') == '' for c in r)]
    
    if len(data_rows) < 1:
        return
    
    num_cols = len(data_rows[0])
    table = doc.add_table(rows=len(data_rows), cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for i, row_data in enumerate(data_rows):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            if j < len(row.cells):
                cell = row.cells[j]
                cell.text = cell_text
                
                # 设置单元格字体
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    for run in paragraph.runs:
                        run.font.name = 'Microsoft YaHei'
                        run.font.size = Pt(10)
                
                # 表头样式
                if i == 0:
                    set_cell_shading(cell, 'CC0000')
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.color.rgb = RGBColor(255, 255, 255)
                else:
                    # 交替行颜色
                    if i % 2 == 0:
                        set_cell_shading(cell, 'F5F5F5')
    
    doc.add_paragraph()


def parse_markdown(md_content):
    """解析Markdown内容"""
    lines = md_content.split('\n')
    doc = create_styled_document()
    
    i = 0
    in_code_block = False
    code_content = []
    
    while i < len(lines):
        line = lines[i]
        
        # 代码块处理
        if line.strip().startswith('```'):
            if in_code_block:
                # 结束代码块
                code_text = '\n'.join(code_content)
                p = doc.add_paragraph()
                run = p.add_run(code_text)
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
                p.paragraph_format.left_indent = Cm(1)
                code_content = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue
        
        if in_code_block:
            code_content.append(line)
            i += 1
            continue
        
        # 标题处理
        if line.startswith('# '):
            text = line[2:].strip()
            doc.add_heading(text, level=1)
        elif line.startswith('## '):
            text = line[3:].strip()
            doc.add_heading(text, level=2)
        elif line.startswith('### '):
            text = line[4:].strip()
            doc.add_heading(text, level=3)
        elif line.startswith('#### '):
            text = line[5:].strip()
            doc.add_heading(text, level=4)
        
        # 分隔线
        elif line.strip() == '---':
            p = doc.add_paragraph()
            p.add_run('─' * 50)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 引用块
        elif line.startswith('> '):
            text = line[2:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            run = p.add_run(text)
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        
        # 表格
        elif '|' in line and i + 1 < len(lines) and '|' in lines[i + 1]:
            rows, new_idx = parse_table(lines, i)
            add_table(doc, rows)
            i = new_idx
            continue
        
        # 无序列表
        elif line.strip().startswith('- '):
            text = line.strip()[2:]
            p = doc.add_paragraph(style='List Bullet')
            process_inline_formatting(p, text)
        
        # 有序列表
        elif re.match(r'^\d+\.\s', line.strip()):
            text = re.sub(r'^\d+\.\s', '', line.strip())
            p = doc.add_paragraph(style='List Number')
            process_inline_formatting(p, text)
        
        # 空行
        elif line.strip() == '':
            pass
        
        # 普通段落
        else:
            text = line.strip()
            if text:
                p = doc.add_paragraph()
                process_inline_formatting(p, text)
        
        i += 1
    
    return doc


def process_inline_formatting(paragraph, text):
    """处理行内格式（粗体、斜体、代码等）"""
    # 处理粗体
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.font.bold = True
        elif part:
            # 处理行内代码
            code_parts = re.split(r'(`[^`]+`)', part)
            for cp in code_parts:
                if cp.startswith('`') and cp.endswith('`'):
                    run = paragraph.add_run(cp[1:-1])
                    run.font.name = 'Consolas'
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
                elif cp:
                    paragraph.add_run(cp)


def convert_md_to_docx(md_file, docx_file):
    """转换Markdown文件为Word文档"""
    print(f"正在读取: {md_file}")
    
    with open(md_file, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    print("正在解析Markdown内容...")
    doc = parse_markdown(md_content)
    
    print(f"正在保存: {docx_file}")
    doc.save(docx_file)
    
    print(f"[OK] 转换完成: {docx_file}")


if __name__ == '__main__':
    import os
    
    # 获取脚本所在目录
    script_dir = os.path.dirname(os.path.abspath(__file__))
    
    # 输入输出文件路径
    md_file = os.path.join(script_dir, '医疗网站内容差距分析.md')
    docx_file = os.path.join(script_dir, '医疗网站内容差距分析.docx')
    
    if os.path.exists(md_file):
        convert_md_to_docx(md_file, docx_file)
    else:
        print(f"[ERROR] 文件不存在: {md_file}")
